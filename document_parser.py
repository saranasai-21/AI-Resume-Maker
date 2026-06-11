import io
import streamlit as st

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

def extract_text_from_file(uploaded_file) -> str:
    """
    Extract text content from an uploaded PDF, DOCX, or text file.
    """
    if not uploaded_file:
        return ""
        
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".pdf"):
        if PdfReader is None:
            st.error("pypdf library is not installed. Cannot parse PDF.")
            return ""
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            st.error(f"Error reading PDF: {e}")
            return ""
            
    elif filename.endswith(".docx"):
        if Document is None:
            st.error("python-docx library is not installed. Cannot parse DOCX.")
            return ""
        try:
            doc = Document(uploaded_file)
            text = ""
            # Paragraph extraction
            for para in doc.paragraphs:
                p_text = para.text.strip()
                if p_text:
                    text += p_text + "\n"
            
            # Table extraction (BUG 2 fix)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading DOCX: {e}")
            return ""
            
    else:
        # Default to reading as plain text UTF-8
        try:
            return uploaded_file.getvalue().decode("utf-8").strip()
        except Exception as e:
            try:
                # Try latin-1 encoding if UTF-8 fails
                return uploaded_file.getvalue().decode("latin-1").strip()
            except Exception as e2:
                st.error(f"Failed to read file as text: {e2}")
                return ""

def convert_html_to_text(html_content: str) -> str:
    """
    Convert HTML resume structure into readable markdown/plain text format.
    """
    import re
    text = html_content
    # Replace headers with markdown equivalent
    text = re.sub(r'</?(h1|H1)[^>]*>', '\n# ', text)
    text = re.sub(r'</?(h2|H2)[^>]*>', '\n\n## ', text)
    text = re.sub(r'</?(h3|H3)[^>]*>', '\n\n### ', text)
    # Replace list tags
    text = re.sub(r'</?li[^>]*>', '\n- ', text)
    # Replace line breaks and paragraph/div containers
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</?(p|div|section)[^>]*>', '\n', text)
    # Strip all other HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Clean up double/triple newlines and spacing
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    return text.strip()

def parse_markdown_to_html(markdown_text: str) -> str:
    """
    Convert tailored plain text markdown back to clean HTML structure locally without APIs.
    """
    import re
    lines = [line.strip() for line in markdown_text.strip().split("\n")]
    html_out = []
    
    # Helper to clean inline formatting (bold/italic)
    def clean_inline(text: str) -> str:
        # Convert **bold** to <strong>bold</strong>
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        # Convert *italic* or _italic_ to <em>italic</em>
        text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.*?)_', r'<em>\1</em>', text)
        return text

    # Extract header info (Name, Subtitle, Contact)
    header_lines = []
    rest_lines = []
    
    # Collect first 3 non-empty lines for the header
    for line in lines:
        if len(header_lines) < 3:
            if line:
                # Remove markdown headers if present
                clean_line = line.lstrip("#").strip()
                header_lines.append(clean_line)
        else:
            rest_lines.append(line)
            
    # Render header
    if len(header_lines) >= 1:
        html_out.append(f"<h1>{clean_inline(header_lines[0])}</h1>")
    if len(header_lines) >= 2:
        html_out.append(f'<div style="text-align: center; font-size: 11px; font-weight: bold; margin-top: -2px; margin-bottom: 6px; color: #4b5563;">{clean_inline(header_lines[1])}</div>')
    if len(header_lines) >= 3:
        html_out.append(f'<div class="contact-info">{clean_inline(header_lines[2])}</div>')
        
    in_list = False
    
    # Process remaining lines
    for line in rest_lines:
        if not line:
            if in_list:
                html_out.append("</ul>")
                in_list = False
            continue
            
        # Check if it is a main header (all caps or ## or #)
        is_main_header = False
        clean_text = line
        if line.startswith("##") or line.startswith("#"):
            is_main_header = True
            clean_text = line.lstrip("#").strip()
        elif line.isupper() and len(line) < 30:
            is_main_header = True
            
        if is_main_header:
            if in_list:
                html_out.append("</ul>")
                in_list = False
            html_out.append(f"<h2>{clean_inline(clean_text)}</h2>")
            continue
            
        # Check if it is a project / experience title (starts with ### or contains " | ", " -- ", or " - ")
        is_sub_header = False
        clean_text = line
        if line.startswith("###"):
            is_sub_header = True
            clean_text = line.lstrip("#").strip()
        elif not (line.startswith("- ") or line.startswith("* ") or line.startswith("+ ")) and (
            " | " in line or " -- " in line or " - " in line or line.startswith("Project:") or line.startswith("Role:")
        ) and len(line) < 120:
            is_sub_header = True
            clean_text = line.lstrip("#").strip()
            
        if is_sub_header:
            if in_list:
                html_out.append("</ul>")
                in_list = False
            html_out.append(f"<h3>{clean_inline(clean_text)}</h3>")
            continue

        # Check if it is a date/location line (often italicized or short metadata like *May 2025 - Jun 2025*)
        is_date_line = False
        if not (line.startswith("- ") or line.startswith("* ") or line.startswith("+ ")) and (
            (line.startswith("*") and line.endswith("*")) or
            any(m in line for m in ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec", "202"])
        ) and len(line) < 60:
            is_date_line = True
            clean_text = line.strip("*").strip()
            
        if is_date_line:
            if in_list:
                html_out.append("</ul>")
                in_list = False
            html_out.append(f'<p style="font-style: italic; margin-top: -3px; margin-bottom: 3px; color: #4b5563; font-size: 9.5px;">{clean_inline(clean_text)}</p>')
            continue
            
        # Check if it is a list item
        if line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
            clean_text = line[2:].strip()
            if not in_list:
                html_out.append("<ul>")
                in_list = True
            html_out.append(f"<li>{clean_inline(clean_text)}</li>")
            continue
            
        # Otherwise it is a normal paragraph
        if in_list:
            html_out.append("</ul>")
            in_list = False
        html_out.append(f"<p>{clean_inline(line)}</p>")
        
    if in_list:
        html_out.append("</ul>")
        
    return "\n".join(html_out)

def generate_pdf_bytes(html_str: str) -> bytes:
    """
    Generate PDF bytes server-side from the HTML string. (BUG 6 fix)
    Uses WeasyPrint.
    """
    try:
        from weasyprint import HTML
        if "<style>" not in html_str:
            import prompts
            html_str = prompts.HTML_PDF_TEMPLATE.format(html_content=html_str)
        return HTML(string=html_str).write_pdf()
    except Exception as e:
        print(f"[ResumeAI] WeasyPrint PDF generation failed: {e}", flush=True)
        return None

def generate_docx_bytes(html_str: str) -> bytes:
    """
    Uses python-docx and bs4 to parse HTML sections into a DOCX document. (BUG 6 fix)
    Returns the DOCX as bytes via io.BytesIO.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from bs4 import BeautifulSoup
    import io
    
    doc = Document()
    
    # Set 0.5 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    soup = BeautifulSoup(html_str, 'html.parser')
    root = soup.body if soup.body else soup
    
    for el in root.children:
        if el.name is None:
            continue
        tag = el.name.lower()
        text = el.get_text().strip()
        if not text:
            continue
            
        if tag == 'h1':
            p = doc.add_paragraph()
            p.alignment = 1 # Center
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            
        elif tag == 'h2':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            
        elif tag == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
        elif tag == 'ul':
            for li in el.find_all('li'):
                li_text = li.get_text().strip()
                if li_text:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(li_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    
        elif tag in ['p', 'div']:
            is_contact = 'contact-info' in el.get('class', []) or 'contact' in text.lower() and len(text) < 150
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            if is_contact:
                p.alignment = 1 # Center
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
            else:
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
